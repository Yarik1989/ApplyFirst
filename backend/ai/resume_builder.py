import copy
import io
import logging
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .tailor import ResumePatch

log = logging.getLogger(__name__)


def _resolve_parent(doc: Any, pointer: str) -> tuple[Any, Any]:
    if not pointer.startswith("/"):
        raise ValueError(f"JSON Pointer must start with '/': {pointer!r}")
    parts = [p.replace("~1", "/").replace("~0", "~") for p in pointer[1:].split("/")]
    *head, last = parts
    node = doc
    for p in head:
        key: Any = int(p) if isinstance(node, list) else p
        node = node[key]
    last_key: Any = int(last) if isinstance(node, list) else last
    return node, last_key


def apply_patches(resume: dict, patches: list[ResumePatch]) -> dict:
    patched = copy.deepcopy(resume)
    for patch in patches:
        try:
            parent, key = _resolve_parent(patched, patch.path)
            current = parent[key]
        except (KeyError, IndexError, ValueError, TypeError) as exc:
            log.warning("Skipping patch %s: resolve failed (%s)", patch.path, exc)
            continue
        if not isinstance(current, str):
            log.warning("Skipping patch %s: target is not a string", patch.path)
            continue
        if current != patch.original:
            log.warning(
                "Skipping patch %s: original mismatch (have %r, patch says %r)",
                patch.path, current, patch.original,
            )
            continue
        parent[key] = patch.replacement
    return patched


def build_pdf(resume: dict) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
    )
    styles = getSampleStyleSheet()
    h_style = ParagraphStyle("h", parent=styles["Heading2"], spaceAfter=4)
    body = styles["BodyText"]

    story: list = []
    story.append(Paragraph(f"<b>{resume.get('name', '')}</b>", styles["Title"]))
    if resume.get("location"):
        story.append(Paragraph(resume["location"], body))
    story.append(Spacer(1, 8))

    if resume.get("licenses"):
        story.append(Paragraph("Licenses &amp; Certifications", h_style))
        story.append(Paragraph(", ".join(resume["licenses"]), body))
        story.append(Spacer(1, 6))

    if resume.get("target_roles"):
        story.append(Paragraph("Target Roles", h_style))
        story.append(Paragraph(", ".join(resume["target_roles"]), body))
        story.append(Spacer(1, 6))

    if resume.get("experience"):
        story.append(Paragraph("Experience", h_style))
        for role in resume["experience"]:
            line = (
                f"<b>{role.get('title', '')}</b> — "
                f"{role.get('org', '')} ({role.get('dates', '')})"
            )
            story.append(Paragraph(line, body))
        story.append(Spacer(1, 6))

    if resume.get("education"):
        story.append(Paragraph("Education", h_style))
        for edu in resume["education"]:
            story.append(Paragraph(edu, body))
        story.append(Spacer(1, 6))

    if resume.get("keywords"):
        story.append(Paragraph("Skills &amp; Keywords", h_style))
        story.append(Paragraph(", ".join(resume["keywords"]), body))

    doc.build(story)
    return buf.getvalue()


def build_docx(resume: dict) -> bytes:
    d = Document()
    d.add_heading(resume.get("name", ""), level=0)
    if resume.get("location"):
        d.add_paragraph(resume["location"])

    if resume.get("licenses"):
        d.add_heading("Licenses & Certifications", level=2)
        d.add_paragraph(", ".join(resume["licenses"]))

    if resume.get("target_roles"):
        d.add_heading("Target Roles", level=2)
        d.add_paragraph(", ".join(resume["target_roles"]))

    if resume.get("experience"):
        d.add_heading("Experience", level=2)
        for role in resume["experience"]:
            p = d.add_paragraph()
            p.add_run(role.get("title", "")).bold = True
            p.add_run(f" — {role.get('org', '')} ({role.get('dates', '')})")

    if resume.get("education"):
        d.add_heading("Education", level=2)
        for edu in resume["education"]:
            d.add_paragraph(edu)

    if resume.get("keywords"):
        d.add_heading("Skills & Keywords", level=2)
        d.add_paragraph(", ".join(resume["keywords"]))

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
